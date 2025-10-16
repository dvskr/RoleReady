from docxtpl import DocxTemplate
from io import BytesIO
from typing import Dict
import os

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "templates", "docx")

def build_docx_template(data: Dict, template_name: str = "classic.docx") -> bytes:
    """Build a DOCX document from a template and data"""
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    
    if not os.path.exists(template_path):
        # Fallback to basic template if file doesn't exist
        return _build_basic_docx(data)
    
    doc = DocxTemplate(template_path)
    doc.render(data)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def _build_basic_docx(data: Dict) -> bytes:
    """Fallback basic DOCX builder if template doesn't exist"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    
    doc = Document()
    
    # Add name as title
    name = data.get('name', 'Resume')
    h = doc.add_heading(name, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary
    summary = data.get('summary', '')
    if summary:
        doc.add_heading('Summary', level=1)
        p = doc.add_paragraph(summary)
        p.style.font.size = Pt(11)
    
    # Add skills
    skills = data.get('skills', [])
    if skills:
        doc.add_heading('Skills', level=1)
        for skill in skills:
            p = doc.add_paragraph(f"• {skill}")
            p.style.font.size = Pt(11)
    
    # Add experience
    experience = data.get('experience', [])
    if experience:
        doc.add_heading('Experience', level=1)
        for job in experience:
            # Job title and company
            title = job.get('title', '')
            company = job.get('company', '')
            dates = job.get('dates', '')
            
            if title or company:
                job_header = f"{title} — {company}"
                if dates:
                    job_header += f" ({dates})"
                p = doc.add_paragraph(job_header)
                p.style.font.size = Pt(11)
                p.style.font.bold = True
            
            # Job bullets
            bullets = job.get('bullets', [])
            for bullet in bullets:
                p = doc.add_paragraph(f"• {bullet}")
                p.style.font.size = Pt(11)
            
            doc.add_paragraph()  # Add space between jobs
    
    # Add education
    education = data.get('education', [])
    if education:
        doc.add_heading('Education', level=1)
        for edu in education:
            degree = edu.get('degree', '')
            school = edu.get('school', '')
            dates = edu.get('dates', '')
            
            if degree or school:
                edu_header = f"{degree} — {school}"
                if dates:
                    edu_header += f" ({dates})"
                p = doc.add_paragraph(edu_header)
                p.style.font.size = Pt(11)
    
    # Save to bytes
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def parse_resume_content(content: str) -> Dict:
    """Parse resume content into structured data for templates"""
    lines = content.split('\n')
    
    data = {
        'name': '',
        'summary': '',
        'skills': [],
        'experience': [],
        'education': []
    }
    
    current_section = None
    current_job = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect sections
        if any(keyword in line.lower() for keyword in ['summary', 'objective', 'profile']):
            current_section = 'summary'
            continue
        elif any(keyword in line.lower() for keyword in ['skills', 'technical skills', 'core strengths']):
            current_section = 'skills'
            continue
        elif any(keyword in line.lower() for keyword in ['experience', 'employment', 'work history']):
            current_section = 'experience'
            continue
        elif any(keyword in line.lower() for keyword in ['education', 'academic']):
            current_section = 'education'
            continue
        
        # Parse content based on current section
        if current_section == 'summary':
            if not data['summary']:
                data['summary'] = line
            else:
                data['summary'] += ' ' + line
        elif current_section == 'skills':
            if line.startswith('•') or line.startswith('-'):
                skill = line[1:].strip()
                if skill:
                    data['skills'].append(skill)
            elif ',' in line:
                # Comma-separated skills
                skills = [s.strip() for s in line.split(',') if s.strip()]
                data['skills'].extend(skills)
        elif current_section == 'experience':
            if line.startswith('•') or line.startswith('-'):
                # Bullet point
                if current_job:
                    bullet = line[1:].strip()
                    if bullet:
                        current_job['bullets'].append(bullet)
            else:
                # Could be job title/company
                if current_job and current_job.get('bullets'):
                    # Previous job is complete
                    data['experience'].append(current_job)
                
                # Start new job
                current_job = {
                    'title': line,
                    'company': '',
                    'dates': '',
                    'bullets': []
                }
        elif current_section == 'education':
            if not data['education']:
                data['education'].append({
                    'degree': line,
                    'school': '',
                    'dates': ''
                })
    
    # Add the last job if exists
    if current_job:
        data['experience'].append(current_job)
    
    # Extract name from first non-empty line if not found
    if not data['name']:
        for line in lines:
            line = line.strip()
            if line and not any(keyword in line.lower() for keyword in ['summary', 'objective', 'experience', 'skills']):
                data['name'] = line
                break
    
    return data
