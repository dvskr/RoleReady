from __future__ import annotations
from io import BytesIO
from typing import Dict, List, Tuple, Optional
import re, difflib
from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph
import fitz  # PyMuPDF
from pdfminer.high_level import extract_text as pdfminer_extract
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image
import numpy as np
import nltk
from nltk.tokenize import sent_tokenize

from .multilingual import multilingual_service

# ---------- Fuzzy heading mapping ----------
CANONICAL = {
    "summary": ["summary","professional summary","profile","about me","overview","executive summary","career summary","objective"],
    "skills": ["skills","technical skills","core competencies","skills & tools","competencies","key skills","technology stack","stack","technical proficiencies","areas of expertise","core skills","skills summary","technology skills","tools & technologies","technical summary","key technologies","tech stack","core strengths"],
    "experience": ["experience","work experience","professional experience","employment","work history","career history","relevant experience"],
    "projects": ["projects","selected projects","personal projects","key projects"],
    "education": ["education","academics","academic background"],
    "certifications": ["certifications","licenses","certs"]
}
MISSPELLINGS = {
    "summery":"professional summary","sumary":"summary","proffesional summary":"professional summary",
    "experiance":"experience","work experiance":"work experience","educatoin":"education","cerifications":"certifications"
}
ALL_ALIASES: Dict[str,str] = {}
for canon,names in CANONICAL.items():
    for n in names: ALL_ALIASES[n]=canon
for wrong,right in MISSPELLINGS.items():
    # map misspelling to same bucket as corrected phrase
    target = None
    for k,v in CANONICAL.items():
        if right in v or right==k:
            target = k; break
    if target: ALL_ALIASES[wrong] = target
ALL_KEYS = list(ALL_ALIASES.keys())

def _norm_heading(s: str) -> str:
    s = s.strip()
    s = re.sub(r"[:\-\u2013\u2014]+$", "", s)
    s = re.sub(r"\s+", " ", s)
    return s.lower()

def _maybe_heading(line: str) -> Optional[str]:
    s = _norm_heading(line)
    if not s or len(s) > 60 or s.endswith("."):
        return None
    if s in ALL_ALIASES:
        return ALL_ALIASES[s]
    match = difflib.get_close_matches(s, ALL_KEYS, n=1, cutoff=0.78)
    if match: return ALL_ALIASES[match[0]]
    if line.isupper():
        match = difflib.get_close_matches(s.title(), ALL_KEYS, n=1, cutoff=0.72)
        if match: return ALL_ALIASES[match[0]]
    return None

# ---------- Cleaners & splitters ----------
_bullet_re = re.compile(r"^(?:[-•·*]|\u2022|\d+\.)\s+")
_space = re.compile(r"\s+")
_newline = re.compile(r"\r\n?")

def _clean_lines(text: str) -> List[str]:
    text = _newline.sub("\n", text)
    lines = [_space.sub(" ", ln).strip() for ln in text.split("\n")]
    return [ln for ln in lines if ln]

def _split_sections(lines: List[str]) -> Dict[str,List[str]]:
    sections: Dict[str,List[str]] = {}
    cur = "other"
    sections[cur] = []
    for ln in lines:
        bucket = _maybe_heading(ln)
        if bucket:
            cur = bucket
            sections.setdefault(cur, [])
        else:
            sections.setdefault(cur, []).append(ln)
    return sections

def _extract_skills(lines: List[str]) -> List[str]:
    text = " ".join(lines)
    
    # Handle comma-separated lists (like "Python, PySpark, SQL, Kafka")
    if "," in text and len(text.split(",")) > 3:
        parts = re.split(r"[,;]|\n", text)
    else:
        parts = re.split(r"[,/|;]|\u2022|\n", text)
    
    skills, seen, out = [], set(), []
    for p in parts:
        t = p.strip()
        # More lenient filtering for skills
        if not t or len(t) > 50 or re.search(r"\d{4}|@|https?://|\.com", t): continue
        # Skip if it's clearly not a skill (dates, locations, etc.)
        if re.search(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b|\b(202[0-9]|202[0-9])\b|\b(Present|St\.|MO|CA|NY|TX|IL|FL)\b", t): continue
        skills.append(t)
    
    for s in skills:
        k = s.lower().strip()
        if k and k not in seen and len(k) > 1:
            seen.add(k); out.append(s)
    return out[:80]

def _extract_bullets_from_lines(lines: List[str]) -> List[str]:
    bullets: List[str] = []
    cur: List[str] = []
    def flush():
        nonlocal bullets, cur
        if cur:
            txt = " ".join(cur).strip()
            if len(txt.split()) >= 5:
                bullets.append(txt)
            cur = []
    for ln in lines:
        if _bullet_re.match(ln):
            flush()
            bullets.append(_bullet_re.sub("", ln).strip()); continue
        if not ln.strip():
            flush(); continue
        # single-sentence heuristic
        if ln.strip().endswith(".") and len(ln.split()) >= 5 and not cur:
            bullets.append(ln.strip())
        else:
            cur.append(ln.strip())
    flush()
    # de-dupe
    out, seen = [], set()
    for b in bullets:
        k = b.lower()
        if k not in seen:
            seen.add(k); out.append(b)
    return out[:200]

# ---------- PDF layout helpers ----------
def _normalize_hyphens(text: str) -> str:
    # join words split across lines: "distribu-\nted" -> "distributed"
    return re.sub(r"(\w)-\n(\w)", r"\1\2", text)

def _page_blocks_to_lines(page: fitz.Page) -> List[str]:
    """
    Use PyMuPDF blocks; keep reading order. Handle simple 2-column pages.
    """
    blocks = page.get_text("blocks")  # each: (x0, y0, x1, y1, text, block_no, ...)
    # filter empties
    blocks = [b for b in blocks if b[4].strip()]
    if not blocks:
        return []

    # Decide if it's likely 2-column by looking at block x positions
    xs = np.array([(b[0]+b[2])/2 for b in blocks])  # centers
    width = page.rect.width
    # fraction of blocks starting in right half
    right_frac = (xs > width * 0.55).mean()
    two_cols = right_frac > 0.3  # heuristic

    lines: List[str] = []

    if not two_cols:
        # simple: sort by y, then x
        for _,_,_,_,text, *_ in sorted(blocks, key=lambda b: (round(b[1],1), round(b[0],1))):
            # block text already multiline; preserve newlines
            text = _normalize_hyphens(text)
            for ln in text.splitlines():
                if ln.strip():
                    lines.append(ln.strip())
        return lines

    # two columns: split blocks by column center, read left column then right
    left_blocks  = [b for b in blocks if (b[0]+b[2])/2 <= width*0.5]
    right_blocks = [b for b in blocks if (b[0]+b[2])/2  > width*0.5]

    for col in (left_blocks, right_blocks):
        for _,_,_,_,text, *_ in sorted(col, key=lambda b: (round(b[1],1), round(b[0],1))):
            text = _normalize_hyphens(text)
            for ln in text.splitlines():
                if ln.strip():
                    lines.append(ln.strip())

    return lines

# ---------- PDF text-layer detection & OCR ----------
def _pdf_text_density(pdf_bytes: bytes, pages_to_check: int = 3) -> float:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    n = min(pages_to_check, len(doc))
    if n == 0: return 0.0
    total_chars = 0
    for i in range(n):
        page = doc.load_page(i)
        txt = page.get_text("text")
        total_chars += len(txt or "")
    doc.close()
    return total_chars / max(n,1)

def _ocr_pdf_to_text(pdf_bytes: bytes, dpi: int = 200) -> str:
    try:
        images = convert_from_bytes(pdf_bytes, dpi=dpi)
        ocr_texts = []
        for img in images:
            if not isinstance(img, Image.Image):
                img = img.convert("RGB")
            ocr_texts.append(pytesseract.image_to_string(img))
        return "\n".join(ocr_texts)
    except Exception as e:
        # Fallback if OCR fails (e.g., Tesseract not installed)
        print(f"OCR failed: {e}")
        return ""

# ---------- DOCX list-aware parsing ----------
def _is_list_paragraph(p: Paragraph) -> bool:
    # True if the paragraph participates in a numbered/bulleted list
    pPr = p._p.pPr
    return (pPr is not None and pPr.numPr is not None)

def _docx_lines_and_lists(doc: Document) -> List[str]:
    """
    Collect lines from paragraphs and tables, preserving list item boundaries.
    We prefix list items so the bullet extractor can split reliably.
    """
    lines: List[str] = []
    def push(text: str, listy: bool):
        t = text.strip()
        if not t: return
        if listy:
            lines.append("• " + t)  # normalize to a bullet marker we recognize
        else:
            lines.append(t)

    # body paragraphs
    for p in doc.paragraphs:
        push(p.text, _is_list_paragraph(p))

    # tables (iterate paragraphs inside cells; preserves bullets too)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    push(p.text, _is_list_paragraph(p))

    return lines

# ---------- DOCX & PDF parsing paths ----------
def _docx_text_with_tables(doc: Document) -> List[str]:
    lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t: lines.append(t)
    return lines

def _parse_from_lines(lines: List[str]) -> Dict:
    lines = _clean_lines("\n".join(lines)) if isinstance(lines, list) else _clean_lines(lines)
    
    # Join all text for language detection
    full_text = " ".join(lines)
    
    # Detect language
    detected_language = multilingual_service.detect_language(full_text)
    
    sections = _split_sections(lines)
    summary = " ".join(
        sections.get("summary", []) or sections.get("professional summary", []) or sections.get("profile", [])
    )[:1000]  # keep it shortish
    skill_lines = (
        sections.get("skills", []) + sections.get("technical skills", []) +
        sections.get("core competencies", []) + sections.get("skills & tools", [])
    )
    skills = _extract_skills(skill_lines)
    
    # Fallback: hunt for category-style lines if skills is empty
    if not skills:
        cat_re = re.compile(r"^(languages|tools|technologies|databases|frameworks|cloud|etl|devops|bi|ml)\s*:\s*(.+)$", re.I)
        extra = []
        for ln in lines if isinstance(lines, list) else lines.split("\n"):
            m = cat_re.match(ln.strip())
            if m:
                items = re.split(r"[,/|;]", m.group(2))
                extra.extend([it.strip() for it in items if it.strip()])
        # dedupe
        seen = set(); out = []
        for s in extra:
            k = s.lower()
            if k and k not in seen and len(s) <= 40:
                seen.add(k); out.append(s)
        skills = out[:80]
    
    exp_lines = (
        sections.get("experience", []) + sections.get("work experience", []) +
        sections.get("professional experience", []) + sections.get("employment", []) +
        sections.get("work history", [])
    )
    experience = _extract_bullets_from_lines(exp_lines) or _extract_bullets_from_lines(lines)
    
    # Add language information to the result
    result = {
        "summary": summary, 
        "skills": skills, 
        "experience": experience,
        "language": detected_language,
        "language_name": multilingual_service.get_language_name(detected_language)
    }
    
    # If not English, translate for embedding purposes and store both versions
    if detected_language != 'en':
        try:
            # Translate summary
            translated_summary = multilingual_service.translate_text(summary, 'en', detected_language) if summary else ""
            
            # Translate skills (limit to avoid API costs)
            translated_skills = []
            for skill in skills[:15]:  # Limit to top 15 skills
                try:
                    translated_skill = multilingual_service.translate_text(skill, 'en', detected_language)
                    translated_skills.append(translated_skill)
                except Exception as e:
                    print(f"Failed to translate skill '{skill}': {e}")
                    translated_skills.append(skill)  # Keep original if translation fails
            
            # Translate experience (limit to avoid API costs)
            translated_experience = []
            for exp in experience[:8]:  # Limit to top 8 experience items
                try:
                    translated_exp = multilingual_service.translate_text(exp, 'en', detected_language)
                    translated_experience.append(translated_exp)
                except Exception as e:
                    print(f"Failed to translate experience '{exp}': {e}")
                    translated_experience.append(exp)  # Keep original if translation fails
            
            # Store translation data
            result.update({
                "translated_content": {
                    "summary": translated_summary,
                    "skills": translated_skills,
                    "experience": translated_experience,
                    "translated_at": datetime.utcnow().isoformat(),
                    "translation_model": "openai-gpt-3.5-turbo"
                }
            })
        except Exception as e:
            print(f"Translation failed: {e}")
            # Store error info but don't fail the entire parsing
            result["translation_error"] = str(e)
    
    return result

def parse_docx(file_bytes: bytes) -> Tuple[Dict,str]:
    doc = Document(BytesIO(file_bytes))
    lines = _docx_lines_and_lists(doc)
    structure = _parse_from_lines(lines)
    return structure, "docx"

def parse_pdf(file_bytes: bytes) -> Tuple[Dict,str]:
    # 1) detect text layer quickly
    density = _pdf_text_density(file_bytes)
    if density < 20:   # lower threshold; some "thin" PDFs still have text
        # OCR fallback
        text = _ocr_pdf_to_text(file_bytes)
        structure = _parse_from_lines(text)
        return structure, "pdf_ocr"

    # 2) Use PyMuPDF blocks for layout-aware extraction
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    all_lines: List[str] = []
    for p in doc:
        all_lines.extend(_page_blocks_to_lines(p))
    doc.close()

    # If PyMuPDF gave almost nothing, try pdfminer as last resort
    if len(all_lines) < 5:
        text = pdfminer_extract(BytesIO(file_bytes)) or ""
        structure = _parse_from_lines(text)
        return structure, "pdf_text"

    # 3) Run our regular splitter/sectioner
    structure = _parse_from_lines(all_lines)

    # 4) If experience is still 0–1 items, sentence-split fallback
    if len(structure.get("experience", [])) <= 1:
        # Find a likely experience chunk from all_lines
        joined = "\n".join(all_lines)
        # Take everything after an experience-ish heading to the next heading or end
        exp_start = re.search(r"(?im)^(experience|work experience|professional experience|employment|work history)[:\s]*$", joined)
        chunk = joined
        if exp_start:
            chunk = joined[exp_start.end():]
        # sentence tokenize (keeps periods)
        try:
            sents = sent_tokenize(chunk)
        except LookupError:
            nltk.download("punkt")
            sents = sent_tokenize(chunk)

        bullets = [s.strip() for s in sents if len(s.split()) >= 6]
        structure["experience"] = bullets[:200]

    return structure, "pdf_text"

def parse_any(filename: str, file_bytes: bytes) -> Dict:
    if filename.lower().endswith(".docx"):
        try:
            # Validate DOCX file by checking ZIP signature
            if len(file_bytes) < 4 or file_bytes[:4] != b'PK\x03\x04':
                raise ValueError("Invalid DOCX file format")
            structure, path = parse_docx(file_bytes)
        except Exception as e:
            # If DOCX parsing fails, fall back to plain text
            print(f"DOCX parsing failed, falling back to plain text: {e}")
            try:
                # Try to decode as text, but filter out binary content
                text = file_bytes.decode(errors="ignore")
                # Remove binary-looking content
                text = ''.join(char for char in text if ord(char) < 128 and char.isprintable() or char.isspace())
                structure = _parse_from_lines(text)
                path = "plain_text_fallback"
            except Exception as e2:
                print(f"Text fallback also failed: {e2}")
                structure = {"summary": "", "skills": [], "experience": [], "confidence": 0.0}
                path = "error"
    elif filename.lower().endswith(".pdf"):
        try:
            structure, path = parse_pdf(file_bytes)
        except Exception as e:
            print(f"PDF parsing failed: {e}")
            structure = {"summary": "", "skills": [], "experience": [], "confidence": 0.0}
            path = "error"
    else:
        # plain text or unknown; try as text
        try:
            text = file_bytes.decode(errors="ignore")
            structure = _parse_from_lines(text)
            path = "plain_text"
        except Exception as e:
            print(f"Text parsing failed: {e}")
            structure = {"summary": "", "skills": [], "experience": [], "confidence": 0.0}
            path = "error"

    # Confidence score: headings found + items count
    conf = 0.0
    if structure.get("summary"): conf += 0.35
    if structure.get("skills"): conf += 0.25
    if len(structure.get("experience", [])) >= 5: conf += 0.25
    if len(structure.get("experience", [])) >= 10: conf += 0.15
    structure["confidence"] = round(conf, 2)
    structure["extraction_path"] = path
    structure["raw_text_present"] = bool(structure.get("summary") or structure.get("experience"))

    return structure

# Public API function for parsing resume content
async def parse_resume_content(text: str) -> Dict:
    """
    Parse resume text into structured sections for the public API
    """
    try:
        lines = text.split('\n')
        structure = _parse_from_lines(lines)
        
        # Add confidence score
        conf = 0.0
        if structure.get("summary"): conf += 0.35
        if structure.get("skills"): conf += 0.25
        if len(structure.get("experience", [])) >= 5: conf += 0.25
        if len(structure.get("experience", [])) >= 10: conf += 0.15
        structure["confidence"] = round(conf, 2)
        
        return {
            "sections": structure,
            "confidence": structure["confidence"]
        }
    except Exception as e:
        return {
            "sections": {
                "personal_info": "",
                "summary": "",
                "experience": [],
                "education": [],
                "skills": []
            },
            "confidence": 0.0,
            "error": str(e)
        }