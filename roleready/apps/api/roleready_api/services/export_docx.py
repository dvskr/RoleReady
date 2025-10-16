from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from io import BytesIO

def build_docx(content: str, title: str = "Resume") -> bytes:
    doc = Document()
    
    # Add title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content with proper formatting
    for line in content.split('\n'):
        if line.strip():  # Skip empty lines
            p = doc.add_paragraph(line.strip())
            p.style.font.size = Pt(11)
    
    # Save to bytes
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
